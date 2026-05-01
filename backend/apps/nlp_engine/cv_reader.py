"""
cv_reader.py
Read raw text from PDF or DOCX CV files.
"""
 
import os
 
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
 
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
 
 
def _read_pdf(file_path: str) -> str:
    """Extract text from a PDF using pdfplumber."""
    if not PDF_AVAILABLE:
        raise ImportError("pdfplumber is not installed. Run: pip install pdfplumber")
 
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"[cv_reader] PDF read error: {e}")
    return text
 
 
def _read_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")
 
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        print(f"[cv_reader] DOCX read error: {e}")
    return text
 
 
def read_cv_file(file_path: str) -> str:
    """
    Read a CV file (PDF or DOCX) and return raw extracted text.
    Raises ValueError for unsupported formats.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
 
    ext = os.path.splitext(file_path)[1].lower()
 
    if ext == ".pdf":
        return _read_pdf(file_path)
    elif ext == ".docx":
        return _read_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF and DOCX are supported.")
 
 
def read_cv_from_django_file(django_file_obj) -> str:
    """
    Read a CV from a Django InMemoryUploadedFile or similar object.
    Saves to a temp file, reads, then deletes it.
    """
    import tempfile
 
    name = django_file_obj.name.lower()
    suffix = ".pdf" if name.endswith(".pdf") else ".docx"
 
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        for chunk in django_file_obj.chunks():
            tmp.write(chunk)
        tmp_path = tmp.name
 
    try:
        text = read_cv_file(tmp_path)
    finally:
        os.unlink(tmp_path)
 
    return text
 